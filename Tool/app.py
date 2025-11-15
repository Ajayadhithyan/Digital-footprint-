from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from pymongo import MongoClient
from bson import ObjectId, Binary
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import hashlib
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import io
import PyPDF2
from docx import Document
import openpyxl
import datetime
import os
import json
import jwt
import secrets
from functools import wraps

# Initialize Flask app
app = Flask(__name__, static_folder='frontend', static_url_path='')
CORS(app)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', secrets.token_hex(32))
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size

# MongoDB Configuration
MONGO_URI = os.getenv('MONGO_URI', 'mongodb://localhost:27017/')
mongodb_connected = False
db = None
files_collection = None
reports_collection = None
users_collection = None
audit_logs = None

def init_mongodb():
    global mongodb_connected, db, files_collection, reports_collection, users_collection, audit_logs
    try:
        client = MongoClient(MONGO_URI, serverSelectionTimeoutMS=5000)
        client.server_info()
        db = client['forensics_toolkit']
        files_collection = db['files']
        reports_collection = db['reports']
        users_collection = db['users']
        audit_logs = db['audit_logs']
        
        # Create indexes for better query performance
        files_collection.create_index([('filename', 'text')])
        files_collection.create_index([('hashes.sha256', 1)])
        files_collection.create_index([('upload_time', -1)])
        users_collection.create_index([('username', 1)], unique=True)
        
        # Create default admin user if not exists
        if users_collection.count_documents({}) == 0:
            default_user = {
                'username': 'admin',
                'password': generate_password_hash('admin123'),
                'email': 'admin@forensics.com',
                'full_name': 'Administrator',
                'organization': 'System Admin',
                'phone': '',
                'role': 'admin',
                'created_at': datetime.datetime.utcnow(),
                'active': True,
                'last_login': None
            }
            users_collection.insert_one(default_user)
            print("[SUCCESS] Default admin user created: admin/admin123")
            
        mongodb_connected = True
        print("[SUCCESS] Successfully connected to MongoDB")
        return True
    except Exception as e:
        print(f"[ERROR] MongoDB Connection Error: {str(e)}")
        mongodb_connected = False
        return False

# Initialize MongoDB connection
init_mongodb()

# JWT Authentication Decorator
def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization')
        if not token:
            return jsonify({'error': 'Token is missing', 'code': 'NO_TOKEN'}), 401
        
        try:
            token = token.split(' ')[1]  # Remove 'Bearer ' prefix
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
            current_user = users_collection.find_one({'username': data['username']})
            if not current_user or not current_user.get('active', True):
                return jsonify({'error': 'Invalid token or user inactive', 'code': 'INVALID_TOKEN'}), 401
        except jwt.ExpiredSignatureError:
            return jsonify({'error': 'Token has expired', 'code': 'TOKEN_EXPIRED'}), 401
        except Exception as e:
            return jsonify({'error': 'Token is invalid', 'code': 'INVALID_TOKEN'}), 401
        
        return f(current_user, *args, **kwargs)
    return decorated

def log_activity(action, user_id=None, file_id=None, details=None):
    """Log all activities for audit trail"""
    if not mongodb_connected:
        return
        
    log_entry = {
        'action': action,
        'user_id': user_id,
        'file_id': file_id,
        'details': details,
        'timestamp': datetime.datetime.utcnow(),
        'ip_address': request.remote_addr,
        'user_agent': request.headers.get('User-Agent', 'Unknown')
    }
    audit_logs.insert_one(log_entry)

# Authentication Routes
@app.route('/api/auth/register', methods=['POST'])
def register():
    try:
        data = request.json
        username = data.get('username', '').strip()
        password = data.get('password', '')
        email = data.get('email', '').strip()
        full_name = data.get('full_name', '').strip()
        organization = data.get('organization', '').strip()
        phone = data.get('phone', '').strip()
        
        # Validate required fields
        if not username or not password:
            return jsonify({'error': 'Username and password required', 'code': 'MISSING_FIELDS'}), 400
        
        if not email:
            return jsonify({'error': 'Email address is required', 'code': 'EMAIL_REQUIRED'}), 400
        
        if len(username) < 3:
            return jsonify({'error': 'Username must be at least 3 characters', 'code': 'USERNAME_TOO_SHORT'}), 400
        
        if len(password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters', 'code': 'PASSWORD_TOO_SHORT'}), 400
        
        # Basic email validation
        if '@' not in email or '.' not in email:
            return jsonify({'error': 'Invalid email address', 'code': 'INVALID_EMAIL'}), 400
        
        # Check if username already exists
        if users_collection.find_one({'username': username}):
            return jsonify({'error': 'Username already exists', 'code': 'USERNAME_EXISTS'}), 400
        
        # Check if email already exists
        if users_collection.find_one({'email': email}):
            return jsonify({'error': 'Email already registered', 'code': 'EMAIL_EXISTS'}), 400
        
        # Create new user with all fields
        user = {
            'username': username,
            'password': generate_password_hash(password),
            'email': email,
            'full_name': full_name if full_name else username,
            'organization': organization if organization else 'Independent',
            'phone': phone if phone else '',
            'role': 'user',
            'created_at': datetime.datetime.utcnow(),
            'active': True,
            'last_login': None
        }
        
        result = users_collection.insert_one(user)
        log_activity('user_registration', user_id=str(result.inserted_id), 
                    details={'username': username, 'email': email, 'organization': organization})
        
        return jsonify({
            'message': 'User created successfully',
            'user': {
                'username': username,
                'email': email,
                'full_name': user['full_name']
            }
        }), 201
        
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'SERVER_ERROR'}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    try:
        data = request.json
        username = data.get('username', '').strip()
        password = data.get('password', '')
        
        if not username or not password:
            return jsonify({'error': 'Username and password required', 'code': 'MISSING_FIELDS'}), 400
        
        user = users_collection.find_one({'username': username})
        if not user or not check_password_hash(user['password'], password):
            return jsonify({'error': 'Invalid credentials', 'code': 'INVALID_CREDENTIALS'}), 401
        
        if not user.get('active', True):
            return jsonify({'error': 'Account deactivated', 'code': 'ACCOUNT_DEACTIVATED'}), 401
        
        # Update last login
        users_collection.update_one(
            {'_id': user['_id']},
            {'$set': {'last_login': datetime.datetime.utcnow()}}
        )
        
        # Generate JWT token
        token = jwt.encode({
            'username': user['username'],
            'role': user['role'],
            'user_id': str(user['_id']),
            'exp': datetime.datetime.utcnow() + datetime.timedelta(hours=24)
        }, app.config['SECRET_KEY'], algorithm='HS256')
        
        log_activity('user_login', user_id=str(user['_id']), details={'username': username})
        
        return jsonify({
            'token': token,
            'user': {
                'username': user['username'],
                'email': user['email'],
                'role': user['role'],
                'user_id': str(user['_id']),
                'full_name': user.get('full_name', user['username']),
                'organization': user.get('organization', 'Independent'),
                'phone': user.get('phone', '')
            }
        }), 200
        
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'SERVER_ERROR'}), 500

@app.route('/api/auth/me', methods=['GET'])
@token_required
def get_current_user(current_user):
    return jsonify({
        'user': {
            'username': current_user['username'],
            'email': current_user['email'],
            'role': current_user['role'],
            'user_id': str(current_user['_id']),
            'created_at': current_user['created_at'].isoformat() if current_user.get('created_at') else None
        }
    }), 200

# Health Check Endpoint
@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "mongodb": "connected" if mongodb_connected else "disconnected",
        "timestamp": datetime.datetime.utcnow().isoformat(),
        "version": "1.0.0"
    })

# Serve Frontend
@app.route('/')
def serve_frontend():
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory(app.static_folder, path)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'pdf', 'docx', 'xlsx', 'pptx', 'txt', 'zip', 'rar'}
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class ForensicAnalyzer:
    @staticmethod
    def generate_hashes(file_data):
        """Generate multiple cryptographic hashes"""
        return {
            'md5': hashlib.md5(file_data).hexdigest(),
            'sha1': hashlib.sha1(file_data).hexdigest(),
            'sha256': hashlib.sha256(file_data).hexdigest(),
            'sha512': hashlib.sha512(file_data).hexdigest()
        }
    
    @staticmethod
    def extract_exif(file_data):
        """Extract comprehensive EXIF data from images"""
        try:
            image = Image.open(io.BytesIO(file_data))
            exif_data = {
                'basic': {
                    'width': image.width,
                    'height': image.height,
                    'format': image.format,
                    'mode': image.mode,
                    'size_bytes': len(file_data)
                },
                'exif': {},
                'gps': None
            }
            
            exif = image._getexif()
            if exif:
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    
                    if tag == 'GPSInfo':
                        gps_data = {}
                        for gps_tag_id, gps_value in value.items():
                            gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                            gps_data[gps_tag] = str(gps_value)
                        exif_data['gps'] = gps_data
                    else:
                        if isinstance(value, bytes):
                            try:
                                value = value.decode('utf-8', errors='ignore')
                            except:
                                value = str(value)
                        exif_data['exif'][tag] = str(value)
            
            return exif_data
        except Exception as e:
            return {'error': str(e), 'basic': {}}
    
    @staticmethod
    def extract_pdf_metadata(file_data):
        """Extract PDF metadata and structure"""
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            metadata = pdf_reader.metadata if hasattr(pdf_reader, 'metadata') and pdf_reader.metadata else {}
            
            pdf_info = {
                'pages': len(pdf_reader.pages),
                'author': metadata.get('/Author', 'Unknown') if metadata else 'Unknown',
                'creator': metadata.get('/Creator', 'Unknown') if metadata else 'Unknown',
                'producer': metadata.get('/Producer', 'Unknown') if metadata else 'Unknown',
                'subject': metadata.get('/Subject', 'Unknown') if metadata else 'Unknown',
                'title': metadata.get('/Title', 'Unknown') if metadata else 'Unknown',
                'created': str(metadata.get('/CreationDate', 'Unknown')) if metadata else 'Unknown',
                'modified': str(metadata.get('/ModDate', 'Unknown')) if metadata else 'Unknown',
                'encrypted': pdf_reader.is_encrypted,
                'page_info': []
            }
            
            for i, page in enumerate(pdf_reader.pages[:3]):
                try:
                    text = page.extract_text()
                    pdf_info['page_info'].append({
                        'page_num': i + 1,
                        'text_length': len(text),
                        'preview': text[:200] if text else 'No text'
                    })
                except:
                    pass
            
            return pdf_info
        except Exception as e:
            return {'error': str(e)}
    
    @staticmethod
    def extract_docx_metadata(file_data):
        """Extract Word document metadata"""
        try:
            doc = Document(io.BytesIO(file_data))
            core_props = doc.core_properties
            
            docx_info = {
                'author': core_props.author or 'Unknown',
                'created': str(core_props.created) if core_props.created else 'Unknown',
                'modified': str(core_props.modified) if core_props.modified else 'Unknown',
                'last_modified_by': core_props.last_modified_by or 'Unknown',
                'title': core_props.title or 'Unknown',
                'subject': core_props.subject or 'Unknown',
                'category': core_props.category or 'Unknown',
                'keywords': core_props.keywords or 'Unknown',
                'comments': core_props.comments or 'None',
                'revision': getattr(core_props, 'revision', 0),
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
            }
            
            text_preview = '\n'.join([p.text for p in doc.paragraphs[:5]])
            docx_info['preview'] = text_preview[:500]
            
            return docx_info
        except Exception as e:
            return {'error': str(e)}
    
    @staticmethod
    def extract_xlsx_metadata(file_data):
        """Extract Excel spreadsheet metadata"""
        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_data))
            
            xlsx_info = {
                'sheets': [sheet.title for sheet in wb.worksheets],
                'sheet_count': len(wb.worksheets),
                'active_sheet': wb.active.title if wb.active else 'None',
                'properties': {}
            }
            
            props = wb.properties
            if props:
                xlsx_info['properties'] = {
                    'creator': props.creator or 'Unknown',
                    'title': props.title or 'Unknown',
                    'subject': props.subject or 'Unknown',
                    'created': str(props.created) if props.created else 'Unknown',
                    'modified': str(props.modified) if props.modified else 'Unknown',
                    'last_modified_by': props.lastModifiedBy or 'Unknown',
                }
            
            first_sheet = wb.worksheets[0]
            xlsx_info['first_sheet_rows'] = first_sheet.max_row
            xlsx_info['first_sheet_cols'] = first_sheet.max_column
            
            return xlsx_info
        except Exception as e:
            return {'error': str(e)}
    
    @staticmethod
    def calculate_security_score(metadata, warnings):
        """Calculate security score based on metadata analysis"""
        score = 100
        
        if warnings:
            score -= len(warnings) * 10
        
        if metadata.get('exif', {}).get('gps'):
            score -= 15
        
        if metadata.get('document', {}).get('author', 'Unknown') != 'Unknown':
            score -= 5
        
        return max(0, min(100, score))

# File Management Routes
@app.route('/api/upload', methods=['POST'])
@token_required
def upload_file(current_user):
    """Upload and analyze file"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided', 'code': 'NO_FILE'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected', 'code': 'NO_FILE_SELECTED'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed', 'code': 'INVALID_FILE_TYPE'}), 400
        
        file_data = file.read()
        if len(file_data) > MAX_FILE_SIZE:
            return jsonify({'error': 'File too large', 'code': 'FILE_TOO_LARGE'}), 413
        
        filename = secure_filename(file.filename)
        file_size = len(file_data)
        file_type = file.content_type or 'application/octet-stream'
        
        # Generate hashes
        hashes = ForensicAnalyzer.generate_hashes(file_data)
        
        # Check for duplicate files
        existing_file = files_collection.find_one({'hashes.sha256': hashes['sha256']})
        if existing_file:
            return jsonify({
                'message': 'File already exists in database',
                'existing_file_id': str(existing_file['_id']),
                'duplicate': True
            }), 200
        
        # Extract metadata based on file type
        metadata = {}
        warnings = []
        
        if file_type.startswith('image/'):
            exif_data = ForensicAnalyzer.extract_exif(file_data)
            metadata['exif'] = exif_data
            
            if exif_data.get('gps'):
                warnings.append('GPS coordinates found in image')
            if exif_data.get('exif', {}).get('Artist') or exif_data.get('exif', {}).get('Copyright'):
                warnings.append('Author/Copyright information present')
            if exif_data.get('exif', {}).get('Software'):
                warnings.append(f"Image edited with: {exif_data['exif']['Software']}")
        
        elif file_type == 'application/pdf' or filename.lower().endswith('.pdf'):
            pdf_meta = ForensicAnalyzer.extract_pdf_metadata(file_data)
            metadata['document'] = pdf_meta
            
            if pdf_meta.get('author') and pdf_meta['author'] != 'Unknown':
                warnings.append(f"Author information present: {pdf_meta['author']}")
            if pdf_meta.get('encrypted'):
                warnings.append('Document is encrypted')
        
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or filename.lower().endswith('.docx'):
            docx_meta = ForensicAnalyzer.extract_docx_metadata(file_data)
            metadata['document'] = docx_meta
            
            if docx_meta.get('author') and docx_meta['author'] != 'Unknown':
                warnings.append(f"Author: {docx_meta['author']}")
            if docx_meta.get('revision') and int(docx_meta.get('revision', 0)) > 0:
                warnings.append(f"Document has {docx_meta['revision']} revisions")
        
        elif file_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' or filename.lower().endswith('.xlsx'):
            xlsx_meta = ForensicAnalyzer.extract_xlsx_metadata(file_data)
            metadata['document'] = xlsx_meta
            
            if xlsx_meta.get('properties', {}).get('creator'):
                warnings.append(f"Creator: {xlsx_meta['properties']['creator']}")
        
        # Calculate security score
        security_score = ForensicAnalyzer.calculate_security_score(metadata, warnings)
        
        # Store file in database
        file_record = {
            'filename': filename,
            'original_filename': file.filename,
            'file_size': file_size,
            'file_type': file_type,
            'hashes': hashes,
            'metadata': metadata,
            'warnings': warnings,
            'security_score': security_score,
            'upload_time': datetime.datetime.utcnow(),
            'uploaded_by': current_user['username'],
            'user_id': str(current_user['_id']),
            'analyzed': True,
            'file_data': Binary(file_data)
        }
        
        result = files_collection.insert_one(file_record)
        file_id = str(result.inserted_id)
        
        # Log activity
        log_activity('file_upload', user_id=str(current_user['_id']), file_id=file_id, 
                    details={'filename': filename, 'size': file_size, 'security_score': security_score})
        
        # Prepare response
        response_data = file_record.copy()
        response_data['_id'] = file_id
        response_data['upload_time'] = response_data['upload_time'].isoformat()
        del response_data['file_data']
        
        return jsonify(response_data), 201
    
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'UPLOAD_ERROR'}), 500

@app.route('/api/files', methods=['GET'])
@token_required
def get_files(current_user):
    """Get all uploaded files with pagination"""
    try:
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 20))
        skip = (page - 1) * limit
        
        # Get filters
        file_type = request.args.get('type')
        has_warnings = request.args.get('warnings', '').lower() == 'true'
        search_query = request.args.get('search', '')
        
        # Build query - filter by current user
        query = {'user_id': str(current_user['_id'])}
        if file_type:
            query['file_type'] = {'$regex': file_type, '$options': 'i'}
        if has_warnings:
            query['warnings'] = {'$exists': True, '$ne': []}
        if search_query:
            query['$or'] = [
                {'filename': {'$regex': search_query, '$options': 'i'}},
                {'original_filename': {'$regex': search_query, '$options': 'i'}}
            ]
        
        # Get total count
        total_count = files_collection.count_documents(query)
        
        # Get files without binary data
        files = list(files_collection.find(
            query,
            {'file_data': 0}
        ).sort('upload_time', -1).skip(skip).limit(limit))
        
        # Convert ObjectId and datetime
        for file in files:
            file['_id'] = str(file['_id'])
            file['upload_time'] = file['upload_time'].isoformat()
        
        return jsonify({
            'files': files,
            'total': total_count,
            'page': page,
            'pages': (total_count + limit - 1) // limit,
            'limit': limit
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'FETCH_ERROR'}), 500

@app.route('/api/files/<file_id>', methods=['GET'])
@token_required
def get_file(current_user, file_id):
    """Get specific file details"""
    try:
        file = files_collection.find_one(
            {'_id': ObjectId(file_id)},
            {'file_data': 0}
        )
        
        if not file:
            return jsonify({'error': 'File not found', 'code': 'FILE_NOT_FOUND'}), 404
        
        # Check if user owns the file or is admin
        if file.get('user_id') != str(current_user['_id']) and current_user.get('role') != 'admin':
            return jsonify({'error': 'Permission denied', 'code': 'PERMISSION_DENIED'}), 403
        
        file['_id'] = str(file['_id'])
        file['upload_time'] = file['upload_time'].isoformat()
        
        log_activity('file_view', user_id=str(current_user['_id']), file_id=file_id, 
                    details={'filename': file['filename']})
        
        return jsonify(file), 200
    
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'FETCH_ERROR'}), 400

@app.route('/api/files/<file_id>/download', methods=['GET'])
@token_required
def download_file(current_user, file_id):
    """Download original file"""
    try:
        file = files_collection.find_one({'_id': ObjectId(file_id)})
        
        if not file:
            return jsonify({'error': 'File not found', 'code': 'FILE_NOT_FOUND'}), 404
        
        # Check if user owns the file or is admin
        if file.get('user_id') != str(current_user['_id']) and current_user.get('role') != 'admin':
            return jsonify({'error': 'Permission denied', 'code': 'PERMISSION_DENIED'}), 403
        
        log_activity('file_download', user_id=str(current_user['_id']), file_id=file_id, 
                    details={'filename': file['filename']})
        
        return send_file(
            io.BytesIO(file['file_data']),
            mimetype=file['file_type'],
            as_attachment=True,
            download_name=file['filename']
        )
    
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'DOWNLOAD_ERROR'}), 400

@app.route('/api/files/<file_id>', methods=['DELETE'])
@token_required
def delete_file(current_user, file_id):
    """Delete a file"""
    try:
        file = files_collection.find_one({'_id': ObjectId(file_id)})
        
        if not file:
            return jsonify({'error': 'File not found', 'code': 'FILE_NOT_FOUND'}), 404
        
        # Check if user owns the file or is admin
        if file.get('user_id') != str(current_user['_id']) and current_user.get('role') != 'admin':
            return jsonify({'error': 'Permission denied', 'code': 'PERMISSION_DENIED'}), 403
        
        result = files_collection.delete_one({'_id': ObjectId(file_id)})
        
        log_activity('file_delete', user_id=str(current_user['_id']), file_id=file_id, 
                    details={'filename': file['filename']})
        
        return jsonify({
            'message': 'File deleted successfully',
            'deleted_count': result.deleted_count
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'DELETE_ERROR'}), 400

# Statistics and Analytics
@app.route('/api/stats', methods=['GET'])
@token_required
def get_stats(current_user):
    """Get comprehensive statistics for current user"""
    try:
        user_id = str(current_user['_id'])
        
        # Filter by current user
        total_files = files_collection.count_documents({'user_id': user_id})
        total_reports = reports_collection.count_documents({'user_id': user_id})
        total_users = users_collection.count_documents({})
        
        # Calculate total storage for user
        pipeline = [
            {'$match': {'user_id': user_id}},
            {'$group': {'_id': None, 'total_size': {'$sum': '$file_size'}}}
        ]
        storage_result = list(files_collection.aggregate(pipeline))
        total_storage = storage_result[0]['total_size'] if storage_result else 0
        
        # Get file type distribution for user
        type_distribution = list(files_collection.aggregate([
            {'$match': {'user_id': user_id}},
            {'$group': {'_id': '$file_type', 'count': {'$sum': 1}, 'total_size': {'$sum': '$file_size'}}}
        ]))
        
        # Get security score distribution for user
        security_distribution = {
            'high': files_collection.count_documents({'user_id': user_id, 'security_score': {'$gte': 80}}),
            'medium': files_collection.count_documents({'user_id': user_id, 'security_score': {'$gte': 50, '$lt': 80}}),
            'low': files_collection.count_documents({'user_id': user_id, 'security_score': {'$lt': 50}})
        }
        
        # Recent activity for user
        recent_files = list(files_collection.find(
            {'user_id': user_id},
            {'filename': 1, 'upload_time': 1, 'security_score': 1, 'file_type': 1}
        ).sort('upload_time', -1).limit(5))
        
        for file in recent_files:
            file['_id'] = str(file['_id'])
            file['upload_time'] = file['upload_time'].isoformat()
        
        # User activity (keep global for display)
        active_users = users_collection.count_documents({'last_login': {'$ne': None}})
        
        return jsonify({
            'total_files': total_files,
            'total_reports': total_reports,
            'total_users': total_users,
            'active_users': active_users,
            'total_storage': total_storage,
            'storage_used_formatted': f"{total_storage / (1024**3):.2f} GB",
            'type_distribution': type_distribution,
            'security_distribution': security_distribution,
            'recent_files': recent_files,
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'STATS_ERROR'}), 500

# Audit Logs
@app.route('/api/audit-logs', methods=['GET'])
@token_required
def get_audit_logs(current_user):
    """Get audit logs for compliance"""
    try:
        if current_user.get('role') != 'admin':
            return jsonify({'error': 'Admin access required', 'code': 'ADMIN_REQUIRED'}), 403
            
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 50))
        skip = (page - 1) * limit
        
        total_count = audit_logs.count_documents({})
        logs = list(audit_logs.find().sort('timestamp', -1).skip(skip).limit(limit))
        
        for log in logs:
            log['_id'] = str(log['_id'])
            log['timestamp'] = log['timestamp'].isoformat()
        
        return jsonify({
            'logs': logs,
            'total': total_count,
            'page': page,
            'pages': (total_count + limit - 1) // limit
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e), 'code': 'LOGS_ERROR'}), 500

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint not found', 'code': 'NOT_FOUND'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error', 'code': 'INTERNAL_ERROR'}), 500

@app.errorhandler(413)
def too_large(error):
    return jsonify({'error': 'File too large', 'code': 'FILE_TOO_LARGE'}), 413

if __name__ == '__main__':
    print("=" * 60)
    print("Digital Forensics Toolkit - Backend Server")
    print("=" * 60)
    print(f"MongoDB: {MONGO_URI}")
    print(f"Server: http://localhost:5000")
    print(f"Database: {db.name if db is not None else 'Not connected'}")
    print(f"Secret Key: {'Set' if app.config['SECRET_KEY'] else 'Not set'}")
    print("=" * 60)
    
    # Create frontend directory if it doesn't exist
    if not os.path.exists('frontend'):
        os.makedirs('frontend')
        print("Created frontend directory")
    
    app.run(host='0.0.0.0', port=5000, debug=True)